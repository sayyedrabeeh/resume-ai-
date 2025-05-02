from django.shortcuts import render
from  docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.http import HttpResponse
# Create your views here.

class GenerateResumeView(APIView):
     permission_classes = [IsAuthenticated]

     def post(self,request):
          data = request.data
          document =Document()

          document.add_heading(data.get_name('name',''),0)

          document.add_paragraph(f"Email: {data.get('email', '')}")
          document.add_paragraph(f"Phone: {data.get('phone', '')}")
          document.add_paragraph(f"Summary:\n{data.get('summary', '')}")

          document.add_heading('Skills',level=1)
          for skill in data.get('skill',[]):
               document.add_paragraph(skill, style='List Bullet')
               
          document.add_heading('Experience',level=1)
          for exp in data.get('Experience',[]):
               document.add_paragraph(f"{exp['title']} at {exp['company']} ({exp['duration']})")
               document.add_paragraph(exp['description'])

          document.add_heading("Education", level=1)
          for edu in data.get("education", []):
            document.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['year']})")


          response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
          response['Content-Disposition'] = 'attachment; filename=resume.docx'
          document.save(response)

          return response
     
